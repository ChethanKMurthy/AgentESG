from __future__ import annotations

import hashlib
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List

from app.core.logging import get_logger

logger = get_logger(__name__)

SUPPORTED_EXTS = {".txt", ".md", ".pdf", ".docx"}


@dataclass
class RawDocument:
    doc_id: str
    source: str
    text: str
    metadata: dict


def _hash(path: str) -> str:
    return hashlib.sha1(path.encode("utf-8")).hexdigest()[:16]


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_docx(path: Path) -> str:
    import docx

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _read_any(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in {".txt", ".md"}:
        return _read_txt(path)
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    raise ValueError(f"Unsupported extension: {ext}")


def iter_paths(root: str) -> Iterable[Path]:
    p = Path(root)
    if p.is_file():
        yield p
        return
    if not p.exists():
        return
    for f in p.rglob("*"):
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTS:
            yield f


def load_documents(root: str) -> List[RawDocument]:
    docs: List[RawDocument] = []
    for path in iter_paths(root):
        try:
            text = _read_any(path)
        except Exception as exc:
            logger.warning("loader_failed", path=str(path), error=str(exc))
            continue
        if not text.strip():
            continue
        docs.append(
            RawDocument(
                doc_id=_hash(str(path)),
                source=str(path.relative_to(Path(root))) if Path(root).is_dir() else path.name,
                text=text,
                metadata={
                    "filename": path.name,
                    "ext": path.suffix.lower(),
                    "size_bytes": os.path.getsize(path),
                },
            )
        )
    logger.info("documents_loaded", count=len(docs), root=root)
    return docs
